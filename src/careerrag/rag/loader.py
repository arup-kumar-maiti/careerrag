"""Load documents and extract structured elements."""

import re
from collections import Counter
from dataclasses import dataclass
from pathlib import Path

import fitz
from docx import Document

BULLET_CHARACTERS = set("●○■□▪▸►•‣◦\u2043\u2022\u2023\u25e6")
BULLET_PATTERN = re.compile(r"^[\u2022\u2023\u25E6\u2043●○■□▪▸►\-\*]\s+")
CONTACT_MAX_LENGTH = 200
CONTACT_PATTERN = re.compile(
    r"[\w.+-]+@[\w-]+\.[\w.]+|https?://\S+|linkedin\.com/\S+|\+?\d[\d\s\-()]{7,}"
)
FONT_SIZE_RATIO_THRESHOLD = 1.15
FONT_SIZE_TITLE_THRESHOLD = 1.4
KIND_BODY = "body"
KIND_CONTACT = "contact"
KIND_HEADING = "heading"
KIND_LIST_ITEM = "list_item"
KIND_SEPARATOR = "separator"
KIND_TABLE = "table"
KIND_TITLE = "title"
MULTIPLE_SPACES = re.compile(r" {2,}")
NOISE_CHARACTERS = re.compile(r"[\xa0\xad\u200b\ufeff]+")
SEPARATOR_PATTERN = re.compile(r"^[\-_=]{3,}\s*$")


@dataclass
class DocumentElement:
    """Represent a structural element of a document."""

    kind: str
    text: str


@dataclass
class LoadedDocument:
    """Represent a parsed document with structured elements."""

    elements: list[DocumentElement]
    source: str


def _classify_line(text: str, font_ratio: float = 1.0) -> str:
    if font_ratio >= FONT_SIZE_TITLE_THRESHOLD:
        return KIND_TITLE
    if font_ratio >= FONT_SIZE_RATIO_THRESHOLD:
        return KIND_HEADING
    if BULLET_PATTERN.match(text):
        return KIND_LIST_ITEM
    if SEPARATOR_PATTERN.match(text):
        return KIND_SEPARATOR
    if CONTACT_PATTERN.match(text) and len(text) < CONTACT_MAX_LENGTH:
        return KIND_CONTACT
    return KIND_BODY


def _classify_docx_paragraph(paragraph: object, text: str) -> str:
    style = getattr(paragraph, "style", None)
    if style and style.name.startswith("Heading"):
        return KIND_HEADING
    return _classify_line(text)


def _load_docx(path: Path) -> list[DocumentElement]:
    doc = Document(str(path))
    elements: list[DocumentElement] = []
    for item in doc.element.body:
        if item.tag.endswith("}p"):
            text = item.text.strip() if hasattr(item, "text") else ""
            if not text:
                continue
            paragraph = next((p for p in doc.paragraphs if p._element is item), None)
            if paragraph:
                kind = _classify_docx_paragraph(paragraph, text)
                elements.append(DocumentElement(kind=kind, text=text))
        elif item.tag.endswith("}tbl"):
            table = next((t for t in doc.tables if t._element is item), None)
            if table:
                rows: list[str] = []
                for row in table.rows:
                    cells = [
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    ]
                    if cells:
                        rows.append(" | ".join(cells))
                if rows:
                    elements.append(
                        DocumentElement(kind=KIND_TABLE, text="\n".join(rows))
                    )
    return elements


def _find_body_font_size(doc: fitz.Document) -> float:
    sizes: Counter[float] = Counter()
    for page in doc:
        for block in page.get_text("dict")["blocks"]:
            if "lines" not in block:
                continue
            for line in block["lines"]:
                for span in line["spans"]:
                    sizes[round(span["size"], 1)] += len(span["text"])
    return sizes.most_common(1)[0][0] if sizes else 12.0


def _find_separator_positions(page: fitz.Page) -> list[float]:
    max_height = 3.0
    min_width_ratio = 0.5
    positions: list[float] = []
    page_width = page.rect.width
    for drawing in page.get_drawings():
        rect = drawing["rect"]
        if rect.height <= max_height and rect.width > page_width * min_width_ratio:
            positions.append(rect.y0)
    return sorted(positions)


def _clean_text(text: str) -> str:
    text = NOISE_CHARACTERS.sub(" ", text)
    return MULTIPLE_SPACES.sub(" ", text).strip()


def _extract_pdf_tables(
    page: fitz.Page,
) -> tuple[list[DocumentElement], list[fitz.Rect]]:
    elements: list[DocumentElement] = []
    table_rects: list[fitz.Rect] = []
    tables = page.find_tables()
    for table in tables.tables:
        table_rects.append(fitz.Rect(table.bbox))
        rows: list[str] = []
        for row in table.extract():
            cells = [_clean_text(cell) for cell in row if cell and cell.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            elements.append(DocumentElement(kind=KIND_TABLE, text="\n".join(rows)))
    return elements, table_rects


def _is_inside_tables(block: dict[str, object], table_rects: list[fitz.Rect]) -> bool:
    block_rect = fitz.Rect(block["bbox"])
    return any(table_rect.intersects(block_rect) for table_rect in table_rects)


def _extract_pdf_page_text(
    page: fitz.Page, body_size: float, table_rects: list[fitz.Rect]
) -> list[DocumentElement]:
    image_block_type = 1
    elements: list[DocumentElement] = []
    for block in page.get_text("dict")["blocks"]:
        if block.get("type") == image_block_type or "lines" not in block:
            continue
        if _is_inside_tables(block, table_rects):
            continue
        for line in block["lines"]:
            line_text = ""
            max_font_size = 0.0
            for span in line["spans"]:
                text = span["text"].strip()
                if not text:
                    continue
                line_text += text + " "
                max_font_size = max(max_font_size, span["size"])
            line_text = _clean_text(line_text)
            if not line_text:
                continue
            font_ratio = max_font_size / body_size if body_size else 1.0
            elements.append(
                DocumentElement(
                    kind=_classify_line(line_text, font_ratio), text=line_text
                )
            )
    return elements


def _merge_bullets(elements: list[DocumentElement]) -> list[DocumentElement]:
    merged: list[DocumentElement] = []
    skip_next = False
    for index, element in enumerate(elements):
        if skip_next:
            skip_next = False
            continue
        if element.text in BULLET_CHARACTERS and index + 1 < len(elements):
            next_element = elements[index + 1]
            merged.append(DocumentElement(kind=KIND_LIST_ITEM, text=next_element.text))
            skip_next = True
        else:
            merged.append(element)
    return merged


def _load_pdf(path: Path) -> list[DocumentElement]:
    doc = fitz.open(path)
    body_size = _find_body_font_size(doc)
    elements: list[DocumentElement] = []
    for page in doc:
        for _ in _find_separator_positions(page):
            elements.append(DocumentElement(kind=KIND_SEPARATOR, text=""))
        table_elements, table_rects = _extract_pdf_tables(page)
        text_elements = _extract_pdf_page_text(page, body_size, table_rects)
        elements.extend(text_elements + table_elements)
    doc.close()
    return _merge_bullets(elements)


def _load_text(path: Path) -> list[DocumentElement]:
    text = path.read_text(encoding="utf-8")
    elements: list[DocumentElement] = []
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        elements.append(DocumentElement(kind=_classify_line(line), text=line))
    return elements


def load_document(path: Path) -> LoadedDocument:
    """Return structured elements and filename from a document path."""
    loaders = {
        ".docx": _load_docx,
        ".md": _load_text,
        ".pdf": _load_pdf,
        ".txt": _load_text,
    }
    elements = loaders[path.suffix.lower()](path)
    return LoadedDocument(elements=elements, source=path.name)
